"""Generador de la proposta d'integració a SAP B1 en format Word (.docx).

Genera `docs/proposta_integracio_sap.docx` per adjuntar al mail al consultor SAP.

Disseny (decisió Oscar 2026-07-24):
- To CONSULTIU/INFORMATIU: no demanem al consultor cap feina tècnica; només li
  presentem el que tenim ara, el que pretenem construir, i li demanem opinió.
- Element mínim a SAP: 3 UDFs a la taula ORDR amb prefix U_FC.
- Càlcul sota demanda: l'usuari marca U_FCCalcular=SÍ i desa; el worker el detecta.
- Sense UDTs, sense User Query, sense panel lateral, sense càlcul automàtic constant.

Ús:
    python scripts/build_proposta_sap.py
"""
from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# Paleta corporativa
COLOR_TITOL = RGBColor(0x2B, 0x6C, 0xB0)
COLOR_SUBTITOL = RGBColor(0x2C, 0x52, 0x82)
COLOR_BODY = RGBColor(0x1A, 0x20, 0x2C)
COLOR_MUTED = RGBColor(0x71, 0x80, 0x8B)
COLOR_TAULA_HDR_BG = "2B6CB0"
COLOR_TAULA_HDR_FG = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_INFO_BG = "E6F0FA"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_TITOL
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_SUBTITOL
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SUBTITOL
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=COLOR_BODY,
             align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.63 + 0.63 * level)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_BODY
    return p


def add_code_block(doc, code, font="Consolas", size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p_pr.append(shd)
    for line in code.rstrip().split("\n"):
        run = p.add_run(line)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = COLOR_BODY
        p.add_run().add_break(WD_BREAK.LINE)
    return p


def add_info_box(doc, text, bg=COLOR_INFO_BG, icon="ℹ"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{icon}  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_BODY
    doc.add_paragraph()


def add_taula(doc, capçaleres, files, amples=None):
    n_cols = len(capçaleres)
    tbl = doc.add_table(rows=1 + len(files), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0]
    for i, cap in enumerate(capçaleres):
        cell = hdr.cells[i]
        set_cell_bg(cell, COLOR_TAULA_HDR_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(cap)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TAULA_HDR_FG
        if amples:
            cell.width = Cm(amples[i])
    for r_idx, fila in enumerate(files):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(fila):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BODY
            if amples:
                cell.width = Cm(amples[c_idx])
    return tbl


def add_page_break(doc):
    doc.add_page_break()


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ====== PORTADA ======
    for _ in range(8):
        doc.add_paragraph()

    add_para(doc, "PROPOSTA D'INTEGRACIÓ", bold=True, size=24,
             color=COLOR_TITOL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Motor d'Embalatges → SAP Business One", bold=True, size=18,
             color=COLOR_SUBTITOL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "Consulta i validació del plantejament", italic=True,
             size=13, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)

    for _ in range(4):
        doc.add_paragraph()

    add_para(doc, f"Data: {date.today().strftime('%d/%m/%Y')}",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Empresa: Agrienergia — Farinera Coromina",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Autor: Oscar Hijazo",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)

    # ====== ÍNDEX ======
    add_heading(doc, "Índex", level=1)
    seccions = [
        ("1", "Què tenim ara: motor d'embalatges"),
        ("2", "Què volem tenir: integració dins SAP"),
        ("3", "Com pensem construir-ho"),
        ("4", "Elements a crear a SAP"),
    ]
    for num, tit in seccions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{num}.  {tit}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_BODY

    add_para(doc, "", space_after=12)

    add_info_box(doc,
        "Aquest document és una consulta prèvia: t'expliquem què hem construït i què "
        "pretenem afegir a SAP, i ens agradaria conèixer la teva opinió abans de tirar-ho "
        "endavant. No et demanem cap feina tècnica — la implementació la farà l'equip intern.")

    add_page_break(doc)

    # ====== SECCIÓ 1 — Què tenim ara ======
    add_heading(doc, "1. Què tenim ara: motor d'embalatges", level=1)
    add_para(doc,
        "El motor d'embalatges és una aplicació web (Python + Flask) que calcula "
        "automàticament els palets d'una comanda de venda a partir de les regles "
        "logístiques del negoci i les dades del client i article.")

    add_heading(doc, "Estat actual", level=2)
    add_bullet(doc, "Variant Kais http://comandes.agrienergia.local/ —, en producció. Llegeix de la BD Kais.")
    add_bullet(doc, "Motor RF1-RF14 idèntic a les dues variants; només canvia la capa de dades.")

    add_heading(doc, "Regles logístiques aplicades (RF1-RF14)", level=2)
    add_taula(doc,
        ["Regla", "Descripció breu"],
        [
            ["RF1", "Filtratge d'articles a granel (GRA) i unitaris (UNI)."],
            ["RF2", "Comanda mínima (40 sacs palet, 20 despaletitzat)."],
            ["RF3", "Comanda mínima de producció per article."],
            ["RF4", "Articles amb dimensió especial."],
            ["RF5", "Prioritat de la direcció d'enviament sobre l'article."],
            ["RF6", "Sac 25 kg especial (tractament diferent d'apilament)."],
            ["RF7", "Aprovisionament d'estoc amb palet específic per producte."],
            ["RF8", "Sac colagne — configuració especial per la família MOULIN DE COLAGNE."],
            ["RF9", "Configuració base per article + direcció + defecte."],
            ["RF10", "Fusió de palets amb mateixa configuració."],
            ["RF11", "Override S05/S10 (max=UxC) si ≥30 sacs."],
            ["RF12", "No barreja d'articles si quantitat múltiple del màxim."],
            ["RF13", "Override de palet client-article (parells específics)."],
            ["RF14", "Fusió de palets residuals amb capacitat lliure."],
        ],
        amples=[1.5, 14.5])

    add_heading(doc, "Com s'usa actualment", level=2)
    add_bullet(doc, "L'operari obre el navegador i accedeix  http://comandes.agrienergia.local/")
    add_bullet(doc, "Cerca la comanda per número.")
    add_bullet(doc, "Veu el resultat calculat (palets, tipus, articles).")
    add_bullet(doc, "Imprimeix l'albarà o exporta el resultat.")

    add_info_box(doc,
        "El pas d'obrir el navegador crea fricció: canvi d'aplicació, dues finestres, "
        "possibilitat d'oblidar-se de recalcular després de modificar la comanda a SAP.")

    add_page_break(doc)

    # ====== SECCIÓ 2 — Què volem tenir ======
    add_heading(doc, "2. Què volem tenir: integració dins SAP", level=1)
    add_para(doc,
        "Volem que el resum del càlcul es vegi directament al formulari Comanda de "
        "venda de SAP, sense que l'operari hagi d'obrir cap altra aplicació. L'objectiu "
        "és eliminar la fricció actual i que la informació sigui sempre coherent amb "
        "l'estat de la comanda quan l'usuari decideix que ja és definitiva.")

    add_heading(doc, "Càlcul sota demanda (no automàtic)", level=2)
    add_para(doc,
        "Sovint les comandes arriben des de la nostra webapp i les usuàries les "
        "modifiquen a SAP afegint o traient sacs iterativament (sac a munt, sac a "
        "vall) fins que queden ajustades. Un càlcul automàtic constant seria molest, "
        "perquè es veurien resultats intermedis que canvien tota l'estona.")
    add_para(doc,
        "Per això proposem un mecanisme senzill: un camp booleà a la comanda que "
        "l'usuari marca quan vol el càlcul. En desar, un worker en background detecta "
        "el flag, fa el càlcul, escriu el resum a la mateixa comanda i torna a posar "
        "el flag a NO. Si més tard cal recalcular, l'usuari torna a marcar el flag.")

    add_heading(doc, "Requisits del disseny", level=2)
    add_bullet(doc, "Compatible amb SAP Web Client i Fat Client (client d'escriptori clàssic).")
    add_bullet(doc, "Càlcul sota demanda: només quan l'usuari el sol·licita amb el flag.")
    add_bullet(doc, "Sense add-ons .NET SDK (no funcionen al Web Client).")
    add_bullet(doc, "Impacte mínim a SAP: només 3 camps addicionals a la taula ORDR. Sense taules noves ni panels laterals.")

    add_heading(doc, "Com es veurà a SAP", level=2)
    add_para(doc,
        "Quan l'operari obre una Comanda de venda a SAP, a la capçalera del formulari "
        "hi haurà 3 camps nous visibles:")
    add_code_block(doc,
        "Comanda: 268/26600093    Client: C211121 LA FLECA DE L'EMPORDA\n"
        "Data:    22/07/2026\n"
        "Import:  ...\n"
        "\n"
        "── Camps nous ─────────────────────────────────────────────\n"
        "Calcular:   ☐  NO   ← l'usuari el marca per sol·licitar el càlcul\n"
        "Embalatge:  3 palets · 120 sacs · palet europeu · CALCULAT\n"
        "Estat:      CALCULAT")

    add_heading(doc, "Flux de treball", level=2)
    add_bullet(doc, "L'operari edita la comanda tantes vegades com calgui, sense pressió de càlculs automàtics.")
    add_bullet(doc, "Quan la comanda ja és definitiva, marca el camp 'Calcular' a SÍ i clica desar.")
    add_bullet(doc, "En pocs segons, el worker calcula i omple els camps 'Embalatge' i 'Estat', i deixa 'Calcular' de nou a NO.")
    add_bullet(doc, "Si posteriorment modifica la comanda, torna a marcar 'Calcular' per obtenir el nou resultat.")

    add_para(doc,
        "Per veure el detall complet dels palets (composició per article, base, apilament), "
        "l'operari pot obrir la web del motor a http://comandes.agrienergia.local/ com sempre ha fet.")

    add_info_box(doc,
        "Aquest disseny és el mínim viable. Si en un futur volguéssim veure també el "
        "detall dels palets dins SAP (sense obrir el navegador), es podria ampliar amb "
        "taules personalitzades (UDTs) i un panel lateral (User-Defined Query). "
        "És un pas opcional que podem fer més endavant si el negoci ho demana.")

    add_page_break(doc)

    # ====== SECCIÓ 3 — Com pensem construir-ho ======
    add_heading(doc, "3. Com pensem construir-ho", level=1)
    add_para(doc,
        "Un procés Python (worker) en background detecta canvis a les comandes de SAP "
        "i escriu el resum als 2 camps nous de la taula ORDR via Service Layer.")

    add_code_block(doc,
        "┌───────────────────────────────────────────────────────────────┐\n"
        "│  SAP Web Client / Fat Client                                  │\n"
        "│  Formulari Sales Order                                        │\n"
        "│    ├─ Camps SAP estàndard (DocEntry, CardCode, Data, ...)     │\n"
        "│    ├─ U_FCCalcular       ← camp nou (SÍ/NO) — flag trigger    │\n"
        "│    ├─ U_FCEmbalatgeResum ← camp nou — resum del càlcul        │\n"
        "│    └─ U_FCEmbalatgeEstat ← camp nou — estat                   │\n"
        "└───────────────────────────────────────────────────────────────┘\n"
        "                              ▲\n"
        "                              │  PATCH via Service Layer\n"
        "                              │  (només quan U_FCCalcular='SÍ')\n"
        "                              │\n"
        "┌─────────────────────────────┴─────────────────────────────────┐\n"
        "│  Host Windows (mateix que ja corre Flask 5002)                │\n"
        "│                                                               │\n"
        "│  ┌─────────────────────┐   ┌─────────────────────────────┐    │\n"
        "│  │  Flask app.py       │   │  sync_worker.py             │    │\n"
        "│  │  Port 5002          │   │  Loop cada 5-10 segons      │    │\n"
        "│  │  (UI web actual —   │   │  ├─ SELECT comandes on      │    │\n"
        "│  │   sense canvis)     │   │  │     U_FCCalcular = 'SÍ'  │    │\n"
        "│  └─────────────────────┘   │  ├─ calcular_embalatges     │    │\n"
        "│                            │  ├─ escriure Resum + Estat  │    │\n"
        "│                            │  └─ posar U_FCCalcular='NO' │    │\n"
        "│                            └─────────────────────────────┘    │\n"
        "└───────────────────────────────────────────────────────────────┘")

    add_heading(doc, "Peces principals", level=2)
    add_bullet(doc, "3 UDFs a la taula ORDR (§4) — visibles a la capçalera de tota Comanda de venda.")
    add_bullet(doc, "1 worker Python que polleja només comandes amb U_FCCalcular='SÍ' (molt eficient).")
    add_bullet(doc, "Cap taula nova, cap User Query, cap panel lateral, cap add-on.")

    add_info_box(doc,
        "Disseny mínim viable amb càlcul sota demanda: 3 camps nous a ORDR i cap alteració "
        "d'estructura de SAP. Molt més eficient que un càlcul automàtic constant, i respecta "
        "el flux natural de les usuàries (editen, ajusten, i només quan estan segures, calculen).")

    add_page_break(doc)

    # ====== SECCIÓ 4 — Elements a crear a SAP ======
    add_heading(doc, "4. Elements a crear a SAP", level=1)
    add_para(doc,
        "Aquesta és la petja completa que la integració deixaria a SAP: 3 camps "
        "addicionals a la taula ORDR. Els crearà l'equip intern com a part del "
        "desenvolupament. Els compartim amb tu perquè validis el disseny.")

    add_heading(doc, "UDFs a la taula ORDR", level=2)
    add_para(doc, "Prefix U_FC coherent amb els UDFs existents (U_FCFabricacionMES, "
             "U_FCVisualitzacioCataleg).")
    add_taula(doc,
        ["UDF", "Tipus", "Mida", "Ubicació", "Descripció"],
        [
            ["U_FCCalcular", "Alfanumèric", "1",
             "Marketing Docs → Title",
             "Flag trigger. Valors: 'S' (calcular) / 'N' (buit) o simple check-box.\nL'usuari el marca a 'S' i desa. El worker el detecta, calcula, i el torna a posar a 'N'."],
            ["U_FCEmbalatgeResum", "Alfanumèric", "254",
             "Marketing Docs → Title",
             "Resum textual del càlcul.\nEx: '3 palets · 120 sacs · palet europeu · CALCULAT'.\nEl worker l'omple."],
            ["U_FCEmbalatgeEstat", "Alfanumèric", "30",
             "Marketing Docs → Title",
             "Estat del càlcul (CALCULAT, CALCULAT_AMB_AVISOS, SOTA_MINIM, NO_CALCULABLE, ERROR).\nPermet filtres i colorejat condicional a llistats."],
        ],
        amples=[3.8, 2.2, 1.2, 3.5, 5.5])

    add_info_box(doc,
        "El worker Python escriurà aquests camps via Service Layer amb PATCH "
        "/Orders({DocEntry}) — no cal cap Formatted Search ni Advanced Layout. "
        "El worker només processa comandes on U_FCCalcular='S', així que la càrrega "
        "és mínima.")

    add_heading(doc, "Alternativa considerada per U_FCCalcular", level=2)
    add_para(doc,
        "Enlloc d'un camp Alfa amb 'S'/'N', podríem usar el tipus 'Valid Values from "
        "List' amb dos valors ('Sí'/'No') o un check-box nadiu, si això s'ajusta millor "
        "a les vostres convencions internes. Ho podem decidir junts.")

    add_para(doc, "Verificació ràpida (per si vols validar via SQL després):", italic=True,
             color=COLOR_MUTED)
    add_code_block(doc,
        "SELECT COLUMN_NAME, DATA_TYPE, CHARACTER_MAXIMUM_LENGTH\n"
        "  FROM INFORMATION_SCHEMA.COLUMNS\n"
        " WHERE TABLE_NAME='ORDR' AND COLUMN_NAME LIKE 'U_FC%Embalatge%' OR COLUMN_NAME='U_FCCalcular'\n"
        "-- Retornarà 3 files: U_FCCalcular nvarchar(1), U_FCEmbalatgeResum nvarchar(254),\n"
        "-- U_FCEmbalatgeEstat nvarchar(30)")

    return doc


def main():
    output_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.abspath(os.path.join(output_dir, "proposta_integracio_sap.docx"))
    doc = build_document()
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"Generat: {output_path} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
