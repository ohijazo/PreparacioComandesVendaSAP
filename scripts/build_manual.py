"""Converteix MANUAL_TECNIC.md a MANUAL_TECNIC.docx i MANUAL_TECNIC.pdf.

Cobreix el subset de Markdown que utilitza el manual:
  - H1 (#), H2 (##), H3 (###)
  - paragraphs
  - llistes amb "- "
  - blocs de codi entre ``` ```
  - taules pipe |...| amb separador |---|
  - emphasis **negreta** i `inline code`
  - separadors horitzontals ---

Ús:
    python scripts/build_manual.py [ruta_md]   # per defecte MANUAL_TECNIC.md

Requereix: python-docx, reportlab. Sense pandoc.
"""
import os
import re
import sys
from dataclasses import dataclass, field

# --- Parser ----------------------------------------------------------------

@dataclass
class Block:
    kind: str               # heading | para | list | code | table | hr
    level: int = 0          # només per heading (1-6)
    text: str = ""          # per para / heading (raw markdown amb inline marks)
    items: list = field(default_factory=list)  # per list
    code: str = ""          # per code (raw)
    lang: str = ""          # per code
    rows: list = field(default_factory=list)   # per table: [[cell,...], ...]
    is_header_row: list = field(default_factory=list)


def parse_md(text: str) -> list[Block]:
    lines = text.split("\n")
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            blocks.append(Block("code", code="\n".join(buf), lang=lang))
            i = j + 1
            continue

        # HR
        if stripped == "---":
            blocks.append(Block("hr"))
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if m:
            blocks.append(Block("heading", level=len(m.group(1)), text=m.group(2)))
            i += 1
            continue

        # Llista
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(Block("list", items=items))
            continue

        # Taula
        if stripped.startswith("|") and stripped.endswith("|"):
            rows = []
            sep_found = False
            while i < len(lines):
                ln = lines[i].strip()
                if not (ln.startswith("|") and ln.endswith("|")):
                    break
                cells = [c.strip() for c in ln.strip("|").split("|")]
                if all(re.match(r"^:?-+:?$", c) for c in cells):
                    sep_found = True
                else:
                    rows.append(cells)
                i += 1
            tb = Block("table", rows=rows)
            tb.is_header_row = [True] + [False] * (len(rows) - 1) if sep_found else [False] * len(rows)
            blocks.append(tb)
            continue

        # Paràgraf (fins línia buida)
        if stripped:
            buf = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip() and not _is_block_start(lines[j].strip()):
                buf.append(lines[j].strip())
                j += 1
            blocks.append(Block("para", text=" ".join(buf)))
            i = j
            continue

        i += 1

    return blocks


def _is_block_start(s: str) -> bool:
    if not s: return True
    if s.startswith("#"): return True
    if s.startswith("```"): return True
    if s.startswith("- ") or s.startswith("* "): return True
    if s.startswith("|"): return True
    if s == "---": return True
    return False


# --- Inline marks: **bold** i `code` ---------------------------------------

_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")

def split_inline(s: str) -> list[tuple[str, str]]:
    """Retorna llista de (tipus, text) on tipus = 'plain' | 'bold' | 'code'."""
    out = []
    last = 0
    for m in _INLINE_RE.finditer(s):
        if m.start() > last:
            out.append(("plain", s[last:m.start()]))
        token = m.group(1)
        if token.startswith("**"):
            out.append(("bold", token[2:-2]))
        else:
            out.append(("code", token[1:-1]))
        last = m.end()
    if last < len(s):
        out.append(("plain", s[last:]))
    return out or [("plain", s)]


# --- DOCX renderer ---------------------------------------------------------

def render_docx(blocks: list[Block], out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Estils base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_inline(p, text):
        for kind, t in split_inline(text):
            run = p.add_run(t)
            if kind == "bold":
                run.bold = True
            elif kind == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)

    for b in blocks:
        if b.kind == "heading":
            level = max(1, min(b.level, 4))
            h = doc.add_heading("", level=level)
            add_inline(h, b.text)
        elif b.kind == "para":
            p = doc.add_paragraph()
            add_inline(p, b.text)
        elif b.kind == "list":
            for it in b.items:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, it)
        elif b.kind == "code":
            # En docx, els \n no es converteixen en salts de línia dins d'un
            # mateix run. Cal afegir-los explícitament amb add_break(LINE).
            from docx.enum.text import WD_BREAK
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            # Manté el bloc de codi sencer a la mateixa pàgina (sense salt
            # de pàgina enmig d'un diagrama o d'una llista de comandes).
            p.paragraph_format.keep_together = True
            lines_code = b.code.split("\n")
            for li, line_code in enumerate(lines_code):
                run = p.add_run(line_code or " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                if li < len(lines_code) - 1:
                    run.add_break(WD_BREAK.LINE)
        elif b.kind == "table":
            if not b.rows:
                continue
            n_cols = max(len(r) for r in b.rows)
            tbl = doc.add_table(rows=len(b.rows), cols=n_cols)
            tbl.style = "Light Grid Accent 1"
            for ri, row in enumerate(b.rows):
                for ci, cell_text in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, cell_text)
                    if b.is_header_row and ri < len(b.is_header_row) and b.is_header_row[ri]:
                        for run in p.runs:
                            run.bold = True
        elif b.kind == "hr":
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(out_path)


# --- PDF renderer (reportlab Platypus) ------------------------------------

def render_pdf(blocks: list[Block], out_path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Preformatted,
        Table, TableStyle, HRFlowable, PageBreak, KeepTogether,
    )
    from reportlab.lib.enums import TA_LEFT

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=1.8*cm, bottomMargin=1.8*cm,
        title="Manual tècnic",
    )
    styles = getSampleStyleSheet()
    h_styles = {
        1: ParagraphStyle("H1", parent=styles["Heading1"], fontSize=20, spaceAfter=10, textColor=colors.HexColor("#1e3a5f")),
        2: ParagraphStyle("H2", parent=styles["Heading2"], fontSize=15, spaceAfter=8, textColor=colors.HexColor("#1e3a5f")),
        3: ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12.5, spaceAfter=6, textColor=colors.HexColor("#2c5282")),
        4: ParagraphStyle("H4", parent=styles["Heading4"], fontSize=11, spaceAfter=4),
    }
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=13, spaceAfter=6)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=20, bulletIndent=8, spaceAfter=2)
    code_style = ParagraphStyle(
        "Code", parent=styles["Code"], fontSize=8.5, leading=10.5,
        backColor=colors.HexColor("#f4f4f4"), borderColor=colors.HexColor("#dddddd"),
        borderWidth=0.5, borderPadding=4, leftIndent=4, rightIndent=4, spaceAfter=8,
    )

    def inline_to_xml(text: str) -> str:
        out = []
        for kind, t in split_inline(text):
            # Escapa XML
            t_esc = (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
            if kind == "bold":
                out.append(f"<b>{t_esc}</b>")
            elif kind == "code":
                out.append(f'<font face="Courier" size="9">{t_esc}</font>')
            else:
                out.append(t_esc)
        return "".join(out)

    flowables = []
    for b in blocks:
        if b.kind == "heading":
            level = max(1, min(b.level, 4))
            flowables.append(Paragraph(inline_to_xml(b.text), h_styles[level]))
        elif b.kind == "para":
            flowables.append(Paragraph(inline_to_xml(b.text), body))
        elif b.kind == "list":
            for it in b.items:
                flowables.append(Paragraph("• " + inline_to_xml(it), bullet))
        elif b.kind == "code":
            # KeepTogether evita que un bloc de codi (p.ex. el diagrama
            # d'arquitectura o una llista de comandes) es talli a meitat
            # de pàgina.
            flowables.append(KeepTogether([Preformatted(b.code, code_style)]))
        elif b.kind == "table" and b.rows:
            data = [[Paragraph(inline_to_xml(c), body) for c in r] for r in b.rows]
            tbl = Table(data, repeatRows=1 if any(b.is_header_row) else 0)
            ts = [
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cccccc")),
                ("LEFTPADDING", (0,0), (-1,-1), 4),
                ("RIGHTPADDING", (0,0), (-1,-1), 4),
                ("TOPPADDING", (0,0), (-1,-1), 3),
                ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ]
            if b.is_header_row and b.is_header_row[0]:
                ts += [
                    ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#e3eaf3")),
                    ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                ]
            tbl.setStyle(TableStyle(ts))
            flowables.append(tbl)
            flowables.append(Spacer(1, 6))
        elif b.kind == "hr":
            flowables.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceBefore=6, spaceAfter=6))

    doc.build(flowables)


# --- Entrypoint ------------------------------------------------------------

def main():
    if len(sys.argv) > 1:
        md_path = sys.argv[1]
    else:
        # Per defecte el manual a l'arrel del repo (un nivell amunt si l'script viu a scripts/)
        here = os.path.dirname(os.path.abspath(__file__))
        md_path = os.path.join(os.path.dirname(here), "MANUAL_TECNIC.md")

    if not os.path.exists(md_path):
        print(f"ERROR: no es troba {md_path}")
        sys.exit(1)

    base, _ = os.path.splitext(md_path)
    docx_path = base + ".docx"
    pdf_path = base + ".pdf"

    with open(md_path, encoding="utf-8") as f:
        text = f.read()

    blocks = parse_md(text)
    print(f"Llegit {md_path}: {len(blocks)} blocs")

    render_docx(blocks, docx_path)
    print(f"Generat {docx_path}")

    render_pdf(blocks, pdf_path)
    print(f"Generat {pdf_path}")


if __name__ == "__main__":
    main()
